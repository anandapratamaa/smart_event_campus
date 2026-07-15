from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

doc = Document()

# Title
title = doc.add_heading('Dokumentasi Sistem Manajemen Event', 0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
subtitle = doc.add_paragraph('Ujian Akhir Semester (UAS)')
subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_heading('1. Pendahuluan', level=1)
doc.add_paragraph('Aplikasi ini adalah Sistem Manajemen Event berbasis Web yang memungkinkan pengguna untuk melakukan pendaftaran pada suatu event. Terdapat dua role utama dalam sistem ini:')
ul = doc.add_paragraph(style='List Bullet')
ul.add_run('Pengguna Umum (User): ').bold = True
ul.add_run('Dapat melihat daftar event yang tersedia dan mendaftarkan diri pada event tersebut.')
ul2 = doc.add_paragraph(style='List Bullet')
ul2.add_run('Administrator (Admin): ').bold = True
ul2.add_run('Memiliki hak akses penuh untuk mengelola (CRUD) data event, melihat daftar pendaftar, serta mengunduh laporan pendaftaran dalam bentuk PDF.')

doc.add_heading('2. Struktur Database', level=1)
doc.add_paragraph('Sistem ini menggunakan database MySQL. Terdapat beberapa tabel utama yang digunakan (berdasarkan database.sql):')
doc.add_paragraph('1. Tabel Users: Menyimpan data admin dan user.', style='List Number')
doc.add_paragraph('2. Tabel Events: Menyimpan informasi terkait acara yang tersedia.', style='List Number')
doc.add_paragraph('3. Tabel Registrations: Menyimpan data peserta yang mendaftar.', style='List Number')

doc.add_heading('3. Penjelasan Fitur dan Antarmuka (Tampilan)', level=1)

# Fitur 1
doc.add_heading('3.1. Halaman Login (login.php)', level=2)
doc.add_paragraph('Fungsi: Memvalidasi kredensial pengguna (username dan password).', style='List Bullet')
doc.add_paragraph('Tampilan: Form login "Secure Login" yang rapi dengan validasi input.', style='List Bullet')
p = doc.add_paragraph('\n[ ---> HAPUS TEKS INI LALU PASTE FOTO LOGIN DI SINI <--- ]\n')
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.runs[0].bold = True

# Fitur 2
doc.add_heading('3.2. Halaman Utama / Beranda (index.php)', level=2)
doc.add_paragraph('Fungsi: Menampilkan daftar event yang tersedia.', style='List Bullet')
doc.add_paragraph('Tampilan: Terdapat daftar event (seperti lomba, seminar) dalam bentuk card.', style='List Bullet')
p = doc.add_paragraph('\n[ ---> HAPUS TEKS INI LALU PASTE FOTO BERANDA DI SINI <--- ]\n')
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.runs[0].bold = True

# Fitur 3
doc.add_heading('3.3. Halaman Dashboard Admin (admin.php)', level=2)
doc.add_paragraph('Fungsi: Menampilkan seluruh statistik event dalam bentuk visual.', style='List Bullet')
doc.add_paragraph('Tampilan: Ringkasan data kegiatan dan grafik distribusi.', style='List Bullet')
p = doc.add_paragraph('\n[ ---> HAPUS TEKS INI LALU PASTE FOTO ADMIN ATAS DI SINI <--- ]\n')
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.runs[0].bold = True
p = doc.add_paragraph('\n[ ---> HAPUS TEKS INI LALU PASTE FOTO ADMIN BAWAH DI SINI <--- ]\n')
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.runs[0].bold = True

# Fitur 4
doc.add_heading('3.4. Halaman Kelola Event (create_event.php & edit_event.php)', level=2)
doc.add_paragraph('Fungsi: Memasukkan rincian event baru atau mengubah yang sudah ada.', style='List Bullet')
doc.add_paragraph('Tampilan: Form input data event yang lengkap.', style='List Bullet')
p = doc.add_paragraph('\n[ ---> HAPUS TEKS INI LALU PASTE FOTO TAMBAH EVENT DI SINI <--- ]\n')
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.runs[0].bold = True
p = doc.add_paragraph('\n[ ---> HAPUS TEKS INI LALU PASTE FOTO EDIT EVENT DI SINI <--- ]\n')
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.runs[0].bold = True

# Fitur 5
doc.add_heading('3.5. Halaman Pendaftaran Event (register_event.php)', level=2)
doc.add_paragraph('Fungsi: Mengisi form pendaftaran acara.', style='List Bullet')
doc.add_paragraph('Tampilan: Form pendaftaran yang interaktif.', style='List Bullet')
p = doc.add_paragraph('\n[ ---> HAPUS TEKS INI LALU PASTE FOTO PENDAFTARAN DI SINI <--- ]\n')
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.runs[0].bold = True

# Fitur 6
doc.add_heading('3.6. Laporan dan Ekspor PDF', level=2)
doc.add_paragraph('Fungsi: Menghasilkan dokumen cetak berisi daftar seluruh kegiatan atau pendaftar.', style='List Bullet')
doc.add_paragraph('Tampilan: Halaman siap cetak.', style='List Bullet')
p = doc.add_paragraph('\n[ ---> HAPUS TEKS INI LALU PASTE FOTO CETAK LAPORAN KEGIATAN DI SINI <--- ]\n')
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.runs[0].bold = True
p = doc.add_paragraph('\n[ ---> HAPUS TEKS INI LALU PASTE FOTO CETAK LAPORAN PENDAFTAR DI SINI <--- ]\n')
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
p.runs[0].bold = True

# Section 4
doc.add_heading('4. Fitur-Fitur Tambahan & Pembaruan (Patch Terbaru)', level=1)
ul4 = doc.add_paragraph(style='List Bullet')
ul4.add_run('Desain Full Dark Mode & Hamburger Menu: ').bold = True
ul4.add_run('Antarmuka pengguna diperbarui secara total menggunakan desain responsif dengan CSS modern (Glassmorphism) dan perbaikan navbar untuk semua ukuran layar.')
ul4_2 = doc.add_paragraph(style='List Bullet')
ul4_2.add_run('Fitur Unduh Materi Acara: ').bold = True
ul4_2.add_run('Pengguna kini dapat mengunduh materi acara langsung dari halaman beranda atau hasil pencarian jika admin telah mengunggah file tersebut.')
ul4_3 = doc.add_paragraph(style='List Bullet')
ul4_3.add_run('Penyempurnaan Tampilan Admin: ').bold = True
ul4_3.add_run('Warna tombol aksi dan judul tabel pada panel pendaftar telah disesuaikan agar memiliki kontras tinggi dan mudah dibaca.')
ul4_4 = doc.add_paragraph(style='List Bullet')
ul4_4.add_run('Perbaikan Bug & Cache Buster: ').bold = True
ul4_4.add_run('Bug saat admin mengubah data acara (ArgumentCountError) telah diperbaiki, serta ditambahkan mekanisme Cache Buster untuk menjamin gaya CSS terbaru selalu dimuat oleh browser.')

doc.add_heading('5. Kesimpulan', level=1)
doc.add_paragraph('Sistem ini dibangun dengan arsitektur PHP native yang terstruktur, memudahkan pengelolaan event mulai dari publikasi hingga pencatatan peserta pendaftaran dengan antarmuka yang sangat modern dan interaktif.')

doc.save('Dokumentasi_Project.docx')
print("Word document generated.")
